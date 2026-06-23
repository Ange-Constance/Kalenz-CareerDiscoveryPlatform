import base64
import os
import re
from urllib.parse import urlparse

import lime
import lime.lime_text
import numpy as np
import pdfplumber
import requests
from docx import Document
from docx.oxml.ns import qn
from transformers import pipeline

_classifier = None

MIN_WORD_COUNT = 50
MAX_GITHUB_REPOS = 10
MAX_README_CHARS = 2000


def _github_token():
    return os.environ.get("GITHUB_TOKEN", "")


CAREER_CLASSES = [
    "Backend Development",
    "Frontend Development",
    "AI/ML Development",
    "UI/UX Design",
    "DevOps/Cloud",
    "Data Engineering",
    "Cybersecurity",
    "Mobile Development",
    "Product/Project Management",
]

SKILL_KEYWORDS = [
    "python", "javascript", "typescript", "react", "node", "java", "sql",
    "docker", "kubernetes", "aws", "git", "html", "css", "figma", "linux",
    "mongodb", "postgresql", "tensorflow", "pytorch", "fastapi", "django",
    "flutter", "kotlin", "swift", "agile", "scrum", "ci/cd", "terraform",
    "networking", "security", "pandas", "spark", "tableau", "excel",
]


def _github_headers():
    headers = {"Accept": "application/vnd.github+json"}
    token = _github_token()
    if token:
        headers["Authorization"] = f"Bearer {token}"
    return headers


def _parse_github_input(username_or_url):
    """Return (username, repo_name or None)."""
    value = (username_or_url or "").strip()
    if not value:
        raise ValueError("GitHub username or URL is required")

    if value.startswith("http"):
        parsed = urlparse(value)
        parts = [p for p in parsed.path.strip("/").split("/") if p]
        if not parts:
            raise ValueError("Invalid GitHub URL")
        username = parts[0]
        repo = parts[1] if len(parts) >= 2 else None
        return username, repo

    if "/" in value:
        username, repo = value.split("/", 1)
        return username.strip(), repo.strip()

    return value, None


def _decode_readme(content, encoding):
    if not content:
        return ""
    if encoding == "base64":
        try:
            return base64.b64decode(content).decode("utf-8", errors="ignore")
        except Exception:
            return ""
    return content


def _fetch_repo_readme(username, repo):
    url = f"https://api.github.com/repos/{username}/{repo}/readme"
    response = requests.get(url, headers=_github_headers(), timeout=30)
    if response.status_code == 404:
        return ""
    response.raise_for_status()
    data = response.json()
    text = _decode_readme(data.get("content"), data.get("encoding"))
    return text[:MAX_README_CHARS]


def _fetch_user_repos(username):
    url = f"https://api.github.com/users/{username}/repos"
    params = {"sort": "updated", "per_page": MAX_GITHUB_REPOS, "type": "owner"}
    response = requests.get(url, headers=_github_headers(), params=params, timeout=30)
    if response.status_code == 404:
        raise ValueError(f"GitHub user '{username}' not found")
    response.raise_for_status()
    return response.json()


def extract_text_from_github(username_or_url):
    """
    Accept username, profile URL, or repo URL.
    Repo URL → that repo's README only.
    Profile/username → up to 10 public repo READMEs.
    """
    username, repo = _parse_github_input(username_or_url)
    chunks = []

    if repo:
        readme = _fetch_repo_readme(username, repo)
        if readme:
            chunks.append(f"Repository: {username}/{repo}\n{readme}")
    else:
        repos = _fetch_user_repos(username)
        for repo_data in repos[:MAX_GITHUB_REPOS]:
            repo_name = repo_data.get("name")
            if not repo_name:
                continue
            readme = _fetch_repo_readme(username, repo_name)
            if readme:
                chunks.append(f"Repository: {username}/{repo_name}\n{readme}")

    combined = "\n\n".join(chunks).strip()
    if not combined:
        raise ValueError(
            "No README content found on GitHub. Try a profile with public repositories."
        )
    return combined


def extract_text_from_certificate(filepath):
    """Extract text from certificate PDF using pdfplumber."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    raise ValueError("Certificate must be a PDF file")


def combine_inputs(cv_text=None, github_text=None, cert_text=None):
    """Combine available inputs into one labeled text block."""
    sections = []
    if cv_text and str(cv_text).strip():
        sections.append(f"[CV/RESUME]\n{cv_text.strip()}")
    if github_text and str(github_text).strip():
        sections.append(f"[GITHUB PROJECTS]\n{github_text.strip()}")
    if cert_text and str(cert_text).strip():
        sections.append(f"[CERTIFICATES]\n{cert_text.strip()}")

    if not sections:
        raise ValueError("At least one input (CV, GitHub, or certificate) is required")

    return "\n\n".join(sections)


def _extract_docx_textboxes(doc):
    chunks = []
    try:
        for element in doc.element.body.iter():
            if element.tag == qn("w:txbxContent"):
                texts = [
                    node.text.strip()
                    for node in element.iter()
                    if node.tag == qn("w:t") and node.text and node.text.strip()
                ]
                if texts:
                    chunks.append(" ".join(texts))
    except Exception:
        pass
    return chunks


def extract_text_from_docx(filepath):
    doc = Document(filepath)
    full_text = []

    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text.strip())

    full_text.extend(_extract_docx_textboxes(doc))
    return " ".join(full_text)


def extract_text_from_pdf(filepath):
    full_text = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text and text.strip():
                full_text.append(text.strip())
                continue
            words = page.extract_words()
            if words:
                page_text = " ".join(w.get("text", "") for w in words if w.get("text"))
                if page_text.strip():
                    full_text.append(page_text.strip())
    return " ".join(full_text)


def validate_extracted_text(text, min_words=None):
    min_words = min_words if min_words is not None else MIN_WORD_COUNT
    if not text or not str(text).strip():
        raise ValueError(
            "Input appears empty or image-based. Please provide text-based content."
        )
    word_count = len(str(text).split())
    if word_count < min_words:
        raise ValueError(
            f"Not enough text to analyze ({word_count} words). "
            "Please provide more detailed content."
        )
    return text


def extract_text_from_file(filepath, strict=True):
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".docx":
        raw = extract_text_from_docx(filepath)
    elif ext == ".pdf":
        raw = extract_text_from_pdf(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    if strict:
        return validate_extracted_text(raw)
    return raw


def extract_key_skills(text, limit=12):
    """Extract likely skills from profile text."""
    if not text:
        return []
    lower = text.lower()
    found = []
    for skill in SKILL_KEYWORDS:
        if skill in lower:
            label = skill.upper() if len(skill) <= 4 else skill.title()
            if skill == "ci/cd":
                label = "CI/CD"
            found.append(label)
    return found[:limit]


def clean_text(text):
    text = str(text)
    text = re.sub(r'http\S+|www\S+', ' ', text)
    text = re.sub(r'!\[.*?\]\(.*?\)', ' ', text)
    text = re.sub(r'\[.*?\]\(.*?\)', ' ', text)
    text = re.sub(r'#{1,6}\s', ' ', text)
    text = re.sub(r'[*_`~>|]', ' ', text)
    text = re.sub(r'<[^>]+>', ' ', text)
    text = re.sub(r'[^a-zA-Z0-9\s\.\,\-]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    # Truncate to 400 words — safe for DistilBERT 512 token limit
    words = text.split()
    if len(words) > 400:
        text = ' '.join(words[:400])
    return text


def load_model():
    global _classifier
    try:
        model_path = os.path.join(
            os.path.dirname(__file__),
            "models",
            "karrelenz-distilbert",
        )
        _classifier = pipeline(
            "text-classification",
            model=model_path,
            top_k=None,
            truncation=True,
            max_length=512,
        )
        print(f"✅ Model loaded from {model_path}")
        return True
    except Exception as e:
        print(f"❌ Model load failed: {e}")
        return False


def is_model_loaded():
    return _classifier is not None


def predict(text):
    global _classifier
    if _classifier is None:
        load_model()

    cleaned = clean_text(text)

    if len(cleaned.split()) < 10:
        raise ValueError(
            "CV appears empty or too short. "
            "Please upload a text-based CV."
        )

    result = _classifier(cleaned)
    scores = sorted(result[0], key=lambda x: -x["score"])

    predicted_career = scores[0]["label"]
    confidence_pct = round(scores[0]["score"] * 100, 1)
    confidence_scores = {
        item["label"]: round(item["score"], 4)
        for item in scores
    }

    return {
        "predicted_career": predicted_career,
        "confidence_scores": confidence_scores,
        "confidence_pct": confidence_pct,
    }


def explain_prediction(text, num_features=8):
    """
    Use LIME to explain why the model predicted a career.
    Returns top words that contributed positively and negatively.
    """
    global _classifier
    if _classifier is None:
        load_model()

    cleaned = clean_text(text)

    # Get class names in alphabetical order (LIME needs consistent order)
    sample_result = _classifier([cleaned[:100]])
    class_names = sorted(
        [r['label'] for r in sample_result[0]]
    )

    def predict_proba(texts):
        """Predict function for LIME — returns probability matrix"""
        results = _classifier(
            texts,
            truncation=True,
            max_length=512
        )
        proba_matrix = []
        for result in results:
            # Sort by label name for consistency
            sorted_result = sorted(result, key=lambda x: x['label'])
            proba_matrix.append([r['score'] for r in sorted_result])
        return np.array(proba_matrix)

    # Create LIME explainer
    explainer = lime.lime_text.LimeTextExplainer(
        class_names=class_names,
        random_state=42
    )

    # Get predicted class index
    prediction = predict(text)
    predicted_career = prediction['predicted_career']
    predicted_idx = class_names.index(predicted_career)

    # Generate explanation
    explanation = explainer.explain_instance(
        cleaned,
        predict_proba,
        num_features=num_features,
        labels=[predicted_idx],
        num_samples=100  # lower = faster, higher = more accurate
    )

    # Extract feature importance
    feature_weights = explanation.as_list(label=predicted_idx)

    # Separate positive and negative contributors
    positive = [
        {"word": word, "weight": round(weight, 4)}
        for word, weight in feature_weights
        if weight > 0
    ]
    negative = [
        {"word": word, "weight": round(abs(weight), 4)}
        for word, weight in feature_weights
        if weight < 0
    ]

    # Sort by weight descending
    positive = sorted(positive, key=lambda x: -x['weight'])
    negative = sorted(negative, key=lambda x: -x['weight'])

    return {
        "predicted_career": predicted_career,
        "supporting_words": positive,    # words FOR this career
        "contradicting_words": negative, # words AGAINST this career
        "top_feature": positive[0]['word'] if positive else None
    }


def check_hf_reachable():
    return is_model_loaded()
